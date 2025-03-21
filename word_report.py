import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

def generate_charts():
    """
    Генерирует тестовые графики (гистограмму и линейный график),
    сохраняет их в файлы .png
    """
    # 1. Пример гистограммы
    data = np.random.randn(1000)
    plt.hist(data, bins=30, color='blue', edgecolor='black')
    plt.title('Пример гистограммы')
    plt.xlabel('Значение')
    plt.ylabel('Частота')
    plt.savefig('histogram.png', dpi=100, bbox_inches='tight')
    plt.close()

    # 2. Пример линейного графика
    x = np.linspace(0, 10, 100)
    y = np.sin(x)
    plt.plot(x, y, label='sin(x)', color='green')
    plt.title('Пример линейного графика')
    plt.xlabel('X')
    plt.ylabel('sin(X)')
    plt.legend()
    plt.savefig('line_chart.png', dpi=100, bbox_inches='tight')
    plt.close()

def create_word_report():
    """
    Создаёт Word-документ (report.docx) с заголовками, параграфами, таблицей
    и вставленными изображениями.
    """
    # Создаём документ
    doc = Document()

    # Заголовок (уровень 0 — самый крупный)
    doc.add_heading('Отчёт о прогнозировании потребления электроэнергии', 0)

    # Краткое вступление
    doc.add_paragraph(
        'В этом отчёте представлены результаты анализа данных о потреблении '
        'электроэнергии и температуре. Ниже показаны основные этапы и выводы.'
    )

    # Подзаголовок 1
    doc.add_heading('1. Исходные данные и методология', level=1)
    doc.add_paragraph(
        'Для анализа были использованы датасеты Load_history.csv и Temperature_history.csv. '
        'Данные были преобразованы из широкого формата в длинный, создан столбец Datetime, '
        'а также выполнена синхронизация нагрузки и температуры по времени.'
    )

    # Подзаголовок 2
    doc.add_heading('2. Результаты EDA (разведочного анализа)', level=1)

    # Вставка изображения (гистограмма)
    doc.add_paragraph('Ниже приведён пример гистограммы:')
    doc.add_picture('histogram.png', width=Inches(4))

    # Вставка второго изображения (линейный график)
    doc.add_paragraph('Пример линейного графика:')
    doc.add_picture('line_chart.png', width=Inches(4))

    # Подзаголовок 3
    doc.add_heading('3. Моделирование и метрики', level=1)
    doc.add_paragraph(
        'Для прогнозирования использовались модели: Линейная регрессия, RandomForestRegressor '
        'и GradientBoostingRegressor. Ниже приведены основные метрики (MSE, R²).'
    )

    # Пример таблицы
    doc.add_paragraph('Таблица 1. Сравнение метрик моделей')
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light List Accent 1'  # или другой стиль

    # Заполняем заголовок таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Модель'
    hdr_cells[1].text = 'MSE'
    hdr_cells[2].text = 'R²'

    # Пример строк с результатами
    data_rows = [
        ('Лин. регрессия', '10000', '0.72'),
        ('Random Forest', '8000', '0.80'),
        ('Gradient Boosting', '7500', '0.82'),
    ]
    for i, row_data in enumerate(data_rows, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]

    # Подзаголовок 4
    doc.add_heading('4. Выводы', level=1)
    doc.add_paragraph(
        'Температура оказывает заметное влияние на потребление электроэнергии. '
        'Наиболее точные результаты показала модель Gradient Boosting, '
        'с MSE около 7500 и R² = 0.82. '
        'Для дальнейшего улучшения точности можно добавить лаговые признаки, '
        'учесть выходные/праздничные дни и другие погодные факторы (влажность, осадки).'
    )

    # Сохраняем документ
    doc.save('report.docx')
    print("Файл report.docx успешно создан.")

def main():
    # Сначала сгенерируем тестовые графики
    generate_charts()

    # Затем сформируем Word-отчёт
    create_word_report()

if __name__ == '__main__':
    main()
